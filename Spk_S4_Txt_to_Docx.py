import os
import re
from collections import defaultdict
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# 设置路径
source_folder = r"D:\备课\02口语话题\已生成答案"
output_path = os.path.join(source_folder, "汇总口语答案.docx")

# 创建文档对象
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)

# 提取 Part x-y 信息（兼容 Part x 直接变 Part x-1）
pattern = re.compile(r'Part\s*(\d+)(?:-(\d+))?', re.IGNORECASE)
parts_dict = defaultdict(dict)  # 结构: parts_dict[x][y] = filename

# 收集文件信息到字典中
for f in os.listdir(source_folder):
    if f.endswith(".txt"):
        match = pattern.search(f)
        if match:
            x = int(match.group(1))
            y = int(match.group(2)) if match.group(2) else 1
            parts_dict[x][y] = f

# 1. 处理 Part 1 全部 y 升序
for y in sorted(parts_dict[1]):
    filename = parts_dict[1][y]
    filepath = os.path.join(source_folder, filename)
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
    cleaned = re.sub(r'[*#`“”]', '', content).strip()
    doc.add_heading(filename.replace(".txt", ""), level=2)
    for para in cleaned.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

# 2. 获取 Part 2 和 Part 3 的所有 y 值，合并去重后排序
all_y = set(parts_dict[2].keys()) | set(parts_dict[3].keys())
for y in sorted(all_y):
    for x in [2, 3]:  # 先写 Part 2-y，再写 Part 3-y
        if y in parts_dict[x]:
            filename = parts_dict[x][y]
            filepath = os.path.join(source_folder, filename)
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
            cleaned = re.sub(r'[*#`“”]', '', content).strip()
            doc.add_heading(filename.replace(".txt", ""), level=2)
            for para in cleaned.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())

# 保存 Word 文档
doc.save(output_path)
print("已生成 Word 文件，路径：", output_path)