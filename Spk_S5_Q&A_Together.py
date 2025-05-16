import os
import re
from docx import Document

# === 路径设置 ===
import argparse

parser = argparse.ArgumentParser()
parser.add_argument("--input", required=True, help="预填txt文件夹")
parser.add_argument("--output", required=True, help="生成答案文件夹")
args = parser.parse_args()

prefill_folder = args.input
output_folder = args.output

origin_docx_path = os.path.join(output_folder, "汇总口语答案.docx")
output_docx_path = os.path.join(output_folder, "汇总口语答案_标题.docx")
prefill_docx_path = os.path.join(output_folder, "预填内容.docx")

# === Step 1: 自动查找符合格式的预填 txt 文件 ===
prefill_txt = None
for fname in os.listdir(prefill_folder):
    if re.match(r"\d{4}_口语话题_预填\.txt", fname):
        prefill_txt = os.path.join(prefill_folder, fname)
        break

if prefill_txt is None:
    raise FileNotFoundError("未找到“xxxx_口语话题_预填.txt”文件")

# === Step 2: 创建预填内容 Word 并写入清洗后的段落 ===
doc = Document()

clean_line = lambda line: re.sub(r'[*#`“”]', '', line.strip())

with open(prefill_txt, "r", encoding="utf-8") as f:
    raw_lines = f.readlines()

lines = [clean_line(line) for line in raw_lines if clean_line(line) != ""]

for line in lines:
    # 去掉项目符号 "- " 以防止被识别为列表
    if line.lstrip().startswith("- "):
        line = line.lstrip()[2:]
    doc.add_paragraph(line)

# === Step 3: 将 Part 1 / 2 的上一段设置为 Heading 3 ===
pattern_part12 = re.compile(r"Part\s*[12](?:-\d+)?", re.IGNORECASE)
paragraphs = doc.paragraphs

for i in range(1, len(paragraphs)):
    text = paragraphs[i].text.strip()
    if pattern_part12.fullmatch(text):
        prev_para = paragraphs[i - 1]
        if prev_para.text.strip():
            prev_para.style = 'Heading 3'

# 保存预填内容 Word 文档
doc.save(prefill_docx_path)
print("已生成清洗后的预填内容 Word：", prefill_docx_path)

# === Step 4: 合并到汇总文档开头 ===
prefill_doc = Document(prefill_docx_path)
main_doc = Document(origin_docx_path)

for element in main_doc.element.body:
    prefill_doc.element.body.append(element)

prefill_doc.save(output_docx_path)
print("已输出合并文档：", output_docx_path)